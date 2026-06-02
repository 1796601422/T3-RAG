from pathlib import Path
import base64

from docx import Document as DocxDocument

from app.prompts.qa import FALLBACK_ANSWER, build_messages
from app.services.document_parser import DocumentParser


def test_markdown_parser_preserves_heading_level_and_lists(tmp_path: Path) -> None:
    file_path = tmp_path / "demo.md"
    file_path.write_text("# Heading\n\n- item one\n- item two\n\nParagraph text.", encoding="utf-8")

    sections = DocumentParser().parse(file_path, "md")

    assert sections[0].block_type == "heading"
    assert sections[0].section_level == 1
    list_sections = [section for section in sections if section.block_type == "list"]
    assert len(list_sections) == 1
    assert "- item one" in list_sections[0].text
    assert "- item two" in list_sections[0].text
    assert any(section.block_type == "paragraph" for section in sections)


def test_markdown_parser_adds_parent_heading_path(tmp_path: Path) -> None:
    file_path = tmp_path / "nested.md"
    file_path.write_text(
        "# Product\n\n## RPA\n\n### Fish Usage\n\nUse the account binding flow.",
        encoding="utf-8",
    )

    sections = DocumentParser().parse(file_path, "md")

    paragraph = next(section for section in sections if section.block_type == "paragraph")
    assert paragraph.section_title == "Product > RPA > Fish Usage"
    assert paragraph.heading_path == ("Product", "RPA", "Fish Usage")


def test_docx_parser_extracts_tables_as_markdown(tmp_path: Path) -> None:
    file_path = tmp_path / "table.docx"
    document = DocxDocument()
    document.add_heading("Pricing", level=1)
    table = document.add_table(rows=3, cols=2)
    rows = [
        ("Item", "Price"),
        ("Basic", "10"),
        ("Pro", "20"),
    ]
    for table_row, values in zip(table.rows, rows, strict=True):
        for cell, value in zip(table_row.cells, values, strict=True):
            cell.text = value
    document.save(file_path)

    sections = DocumentParser().parse(file_path, "docx")

    table_sections = [section for section in sections if section.block_type == "table"]
    assert len(table_sections) == 1
    assert table_sections[0].section_title == "Pricing"
    assert "| Item | Price |" in table_sections[0].text
    assert "| Basic | 10 |" in table_sections[0].text


def test_docx_parser_extracts_images_with_heading_path(tmp_path: Path) -> None:
    file_path = tmp_path / "image.docx"
    image_path = tmp_path / "flow.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAFgwJ/l4wH4wAAAABJRU5ErkJggg=="
        )
    )
    document = DocxDocument()
    document.add_heading("流程设计", level=1)
    document.add_picture(str(image_path))
    document.save(file_path)

    sections = DocumentParser().parse(file_path, "docx", document_id="doc-image-test")

    image_sections = [section for section in sections if section.block_type == "image"]
    assert len(image_sections) == 1
    assert image_sections[0].section_title == "流程设计"
    assert image_sections[0].heading_path == ("流程设计",)
    assert image_sections[0].image_path
    assert "所属标题：流程设计" in image_sections[0].text


def test_docx_parser_detects_textual_product_headings(tmp_path: Path) -> None:
    file_path = tmp_path / "strategy.docx"
    document = DocxDocument()
    document.add_heading("【智能招募】AI外呼策略优化", level=1)
    document.add_paragraph("版本信息")
    document.add_paragraph("meego：[需求]AI电话-新线索选取优化及策略配置优化")
    document.add_paragraph("一、产品概述")
    document.add_paragraph("1.1 产品背景")
    document.add_paragraph("AI电话邀约和清洗已经上线，经过测试整体效果良好。")
    document.add_paragraph("1.2 产品目标")
    document.add_paragraph("ai外呼-线索意愿清洗场景赋能给自营交付城市销售使用。")
    document.add_paragraph("1.3产品功能列表")
    document.save(file_path)

    sections = DocumentParser().parse(file_path, "docx")

    headings = [section for section in sections if section.block_type == "heading"]
    assert [heading.text for heading in headings] == [
        "【智能招募】AI外呼策略优化",
        "一、产品概述",
        "1.1 产品背景",
        "1.2 产品目标",
        "1.3产品功能列表",
    ]
    assert [heading.section_level for heading in headings] == [1, 1, 2, 2, 2]

    background = next(section for section in sections if "测试整体效果良好" in section.text)
    assert background.section_title == "一、产品概述 > 1.1 产品背景"


def test_docx_parser_does_not_treat_colon_labels_as_headings(tmp_path: Path) -> None:
    file_path = tmp_path / "labels.docx"
    document = DocxDocument()
    document.add_paragraph("一、主要使用场景")
    document.add_paragraph("输出：")
    document.add_paragraph("生成一份调研报告。")
    document.save(file_path)

    sections = DocumentParser().parse(file_path, "docx")

    headings = [section.text for section in sections if section.block_type == "heading"]
    assert headings == ["一、主要使用场景"]
    paragraph = next(section for section in sections if "输出：" in section.text)
    assert paragraph.block_type == "paragraph"
    assert paragraph.section_title == "一、主要使用场景"


def test_prompt_build_messages_includes_context_and_fallback() -> None:
    messages = build_messages(
        "What happened?",
        [
            {
                "filename": "demo.pdf",
                "page_no": 2,
                "section_title": "Summary",
                "content": "Important evidence.",
                "retrieval_note": "vector retrieval hit",
            }
        ],
    )

    assert len(messages) == 2
    assert "demo.pdf" in messages[1]["content"]
    assert FALLBACK_ANSWER in messages[1]["content"]
